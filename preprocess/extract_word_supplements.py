#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word supplementary material extractor (Llama 사용 안함)
Extracts compounds from Word files and saves to supplement_extracted/
- 테이블만 추출 (Llama 사용 안함)
- 테이블이 없으면 텍스트만 저장

Usage:
  python extract_word_supplements.py Supplements/PMC###/file.docx
"""

import sys
import json
import pandas as pd
import re
from pathlib import Path
from datetime import datetime
from docx import Document

def extract_tables_from_word(file_path: Path):
    """Extract all tables from Word file"""
    try:
        doc = Document(file_path)
        all_data = []
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            
            if rows:
                all_data.append({
                    'table_index': table_idx,
                    'rows': rows,
                    'max_cols': max(len(r) for r in rows) if rows else 0
                })
        
        return all_data
    except Exception as e:
        print(f"Error reading Word file: {e}")
        return []


def extract_text_from_word(file_path: Path):
    """Extract all text content from Word file"""
    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return '\n'.join(paragraphs)
    except Exception as e:
        print(f"Error extracting text from Word: {e}")
        return ""


def table_to_long_format(table_data, sheet_name="Word Document"):
    """Convert table to long format (compound × attribute × value)"""
    rows = table_data.get('rows', [])
    if not rows:
        return []
    
    # Find header row (first non-empty row)
    header_row_idx = None
    for i, row in enumerate(rows):
        if any(cell.strip() for cell in row):
            header_row_idx = i
            break
    
    if header_row_idx is None:
        return []
    
    headers = rows[header_row_idx]
    
    # Find compound name column (look for patterns like "Compound", "Name", "ID", "SMILES", etc.)
    name_col_idx = None
    name_patterns = [
        r'compound', r'name', r'id', r'smiles', r'molecule', 
        r'chemical', r'drug', r'cbk', r'well', r'position'
    ]
    
    for i, header in enumerate(headers):
        header_lower = str(header).lower()
        if any(re.search(p, header_lower) for p in name_patterns):
            name_col_idx = i
            break
    
    # If no name column found, use first column
    if name_col_idx is None:
        name_col_idx = 0
    
    rows_data = []
    
    # Process data rows (skip header row)
    for row_idx, row in enumerate(rows[header_row_idx + 1:], start=header_row_idx + 1):
        if not any(cell.strip() for cell in row):
            continue
        
        # Get compound name
        compound_name = ""
        if name_col_idx < len(row):
            compound_name = str(row[name_col_idx]).strip()
        
        if not compound_name:
            continue
        
        # Extract attributes (other columns)
        for col_idx, header in enumerate(headers):
            if col_idx == name_col_idx:
                continue
            
            if col_idx >= len(row):
                continue
            
            value = str(row[col_idx]).strip()
            if not value or value.lower() in ['na', 'n/a', '-', '']:
                continue
            
            attribute_name = str(header).strip() if header else f"Column_{col_idx}"
            
            rows_data.append({
                "sheet": sheet_name,
                "source_cell": f"Table{table_data['table_index']}_R{row_idx}_C{col_idx}",
                "row_excel": row_idx,
                "compound_name": compound_name,
                "attribute_name": attribute_name,
                "value": value,
                "guessed_metric": "",  # Can be enhanced later
            })
    
    return rows_data

def main():
    import argparse
    
    parser = argparse.ArgumentParser(description="Extract compounds from Word supplementary files (Llama 사용 안함)")
    parser.add_argument("word_file", help="Path to Word file (.docx)")
    parser.add_argument("--output_dir", "-o", help="Output directory (default: supplement_extracted/PMC###/)")
    
    args = parser.parse_args()
    
    word_path = Path(args.word_file)
    
    if not word_path.exists():
        print(f"[ERR] Word file not found: {word_path}")
        sys.exit(1)
    
    # Auto-detect PMC ID
    pmc_id = None
    for part in word_path.parts:
        if part.startswith("PMC"):
            pmc_id = part
            break
    
    if not pmc_id:
        print("[ERR] PMC ID not found in path")
        sys.exit(1)
    
    # Determine output directory
    if args.output_dir:
        out_dir = Path(args.output_dir)
    else:
        # Default: supplement_extracted/PMC###/
        out_dir = Path("supplement_extracted") / pmc_id
    
    out_dir.mkdir(parents=True, exist_ok=True)
    
    # Extract tables from Word
    print(f"Extracting tables from Word (Llama 사용 X)...")
    tables = extract_tables_from_word(word_path)
    
    # 테이블 데이터를 long format으로 변환
    all_rows = []
    for table_data in tables:
        rows = table_to_long_format(table_data, sheet_name=word_path.stem)
        all_rows.extend(rows)
    
    if all_rows:
        # 테이블 데이터 저장
        df = pd.DataFrame(all_rows)
        
        csv_path = out_dir / f"{pmc_id}_compounds_from_word.csv"
        json_path = out_dir / f"{pmc_id}_compounds_from_word.json"
        
        df.to_csv(csv_path, index=False, encoding="utf-8-sig")
        df.to_json(json_path, orient="records", force_ascii=False, indent=2)
        
        n_compounds = df["compound_name"].nunique()
        n_rows = len(df)
        
        print(f"Word 테이블 추출 완료 (Llama 사용 X)")
        print(f"  word: {word_path}")
        print(f"  compounds: {n_compounds} | rows: {n_rows}")
        print(f"  saved: {csv_path}")
        print(f"  saved: {json_path}")
    else:
        # 테이블이 없으면 텍스트만 추출
        print("테이블 없음, 텍스트만 저장")
        text_content = extract_text_from_word(word_path)
        
        if text_content:
            text_path = out_dir / f"{pmc_id}_from_word.txt"
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            print(f"  텍스트 저장: {text_path}")
        else:
            print(f"텍스트 없음")

if __name__ == "__main__":
    main()


